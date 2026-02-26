import streamlit as st
import asyncio
import edge_tts
import tempfile
import os
import re
import requests
from PyPDF2 import PdfReader
from docx import Document

# Configuración de la página
st.set_page_config(page_title="Beta Lector", page_icon="🔊")
st.title("🔊 Me encargaré de leer tu texto")

# Función para limpiar el texto extraído (eliminar saltos de línea excesivos)
def clean_text(text):
    """
    Limpia el texto para que sea más fluido al leerlo con TTS:
    - Reemplaza saltos de línea simples por espacios.
    - Conserva los dobles saltos como separadores de párrafo.
    - Elimina espacios múltiples.
    """
    # Normalizar saltos de línea: convertir \r\n a \n
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # Reemplazar saltos de línea simples (no consecutivos) por espacios
    # Usamos una expresión regular que busca un salto de línea que no esté seguido de otro salto
    # Pero es más sencillo: primero dividimos por dobles saltos para conservar párrafos
    paragraphs = re.split(r'\n\s*\n', text)
    cleaned_paragraphs = []
    for para in paragraphs:
        # Dentro de cada párrafo, reemplazar saltos simples por espacios
        para = re.sub(r'\n+', ' ', para)
        # Eliminar espacios múltiples
        para = re.sub(r'\s+', ' ', para)
        # Quitar espacios al inicio y final
        para = para.strip()
        if para:
            cleaned_paragraphs.append(para)
    # Unir párrafos con doble salto de línea
    return '\n\n'.join(cleaned_paragraphs)

# Obtener lista de voces disponibles (en caché)
@st.cache_resource
def get_voices():
    voices = asyncio.run(edge_tts.list_voices())
    voice_list = []
    for v in voices:
        friendly_name = f"{v['FriendlyName']} ({v['Gender']}, {v['Locale']})"
        voice_list.append({
            "name": v["Name"],
            "friendly": friendly_name,
            "locale": v["Locale"],
            "gender": v["Gender"]
        })
    return voice_list

voices_data = get_voices()

# Interfaz principal
st.markdown("### 📝 Introduce el texto")

# Opciones de entrada: archivo, Drive o texto manual
opcion_entrada = st.radio(
    "¿Cómo quieres proporcionar el texto?",
    ("Escribir manualmente", "Subir archivo (PDF o DOCX)", "Enlace de Google Drive")
)

texto = ""  # Variable para almacenar el texto final

if opcion_entrada == "Escribir manualmente":
    texto = st.text_area("Texto", height=500, placeholder="Escribe aquí...")

elif opcion_entrada == "Subir archivo (PDF o DOCX)":
    archivo_subido = st.file_uploader("Selecciona un archivo", type=["pdf", "docx"])
    if archivo_subido is not None:
        with st.spinner("Extrayendo y limpiando texto..."):
            texto_extraido = ""
            if archivo_subido.type == "application/pdf":
                pdf_reader = PdfReader(archivo_subido)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        texto_extraido += page_text + "\n"
            elif archivo_subido.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(archivo_subido)
                for para in doc.paragraphs:
                    texto_extraido += para.text + "\n"
            # Limpiar el texto extraído
            texto = clean_text(texto_extraido)
        # Mostrar el texto limpio en un área editable
        texto = st.text_area("Texto extraído (puedes editarlo)", value=texto, height=500)
    else:
        st.info("Por favor, sube un archivo.")

elif opcion_entrada == "Enlace de Google Drive":
    enlace = st.text_input("Pega el enlace público de Google Drive")
    if enlace:
        # Extraer ID del archivo de un enlace de Drive
        match = re.search(r'/d/([a-zA-Z0-9_-]+)', enlace)
        if match:
            file_id = match.group(1)
            download_url = f"https://drive.google.com/uc?export=download&id={file_id}"
            with st.spinner("Descargando y extrayendo texto..."):
                texto_extraido = ""
                try:
                    response = requests.get(download_url)
                    if response.status_code == 200:
                        with tempfile.NamedTemporaryFile(delete=False) as tmp:
                            tmp.write(response.content)
                            tmp_path = tmp.name
                        # Intentar como PDF
                        try:
                            pdf_reader = PdfReader(tmp_path)
                            for page in pdf_reader.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    texto_extraido += page_text + "\n"
                        except:
                            # Intentar como DOCX
                            try:
                                doc = Document(tmp_path)
                                for para in doc.paragraphs:
                                    texto_extraido += para.text + "\n"
                            except:
                                st.error("No se pudo extraer texto del archivo. Asegúrate de que sea un PDF o DOCX válido.")
                        os.unlink(tmp_path)
                    else:
                        st.error("Error al descargar el archivo. Verifica que el enlace sea público.")
                except Exception as e:
                    st.error(f"Error: {e}")
                if texto_extraido:
                    texto = clean_text(texto_extraido)
                    texto = st.text_area("Texto extraído (puedes editarlo)", value=texto, height=500)
        else:
            st.error("No se pudo extraer el ID del archivo del enlace. Asegúrate de usar un enlace de Drive válido (ej. https://drive.google.com/file/d/.../view).")

# ---- Sección de selección de voz ----
st.markdown("### 🎤 ¿Quién te leerá hoy?")

# Filtrar voces por idioma español (opcional, pero podemos mostrar solo español)
idiomas_espanol = [loc for loc in set(v["locale"] for v in voices_data) if loc.startswith("es")]
idioma_sel = st.selectbox("Filtrar por idioma (opcional)", ["Todos"] + sorted(idiomas_espanol))

if idioma_sel == "Todos":
    voces_mostrar = voices_data
else:
    voces_mostrar = [v for v in voices_data if v["locale"] == idioma_sel]

# Crear opciones para el selector de voz
voz_opciones = {v["friendly"]: v["name"] for v in voces_mostrar}

# Preseleccionar español de México, mujer (es-MX-DaliaNeural)
default_voice_name = "es-MX-DaliaNeural"
default_friendly = None
for friendly, name in voz_opciones.items():
    if name == default_voice_name:
        default_friendly = friendly
        break

if default_friendly is None:
    for v in voices_data:
        if v["name"] == default_voice_name:
            default_friendly = v["friendly"]
            if idioma_sel != "Todos" and v["locale"] != idioma_sel:
                st.warning("La voz preseleccionada (es-MX-DaliaNeural) no está en el filtro actual. Se mostrará igualmente.")
            voz_opciones[default_friendly] = default_voice_name
            break

indice_por_defecto = list(voz_opciones.keys()).index(default_friendly) if default_friendly else 0
voz_elegida_nombre = st.selectbox(
    "Selecciona una voz",
    options=list(voz_opciones.keys()),
    index=indice_por_defecto
)
# Botón para generar el audio
if st.button("🔊 Generar audio"):
    if not texto.strip():
        st.warning("Por favor, escribe o proporciona algún texto.")
    else:
        with st.spinner("Generando audio..."):
            voice_name = voz_opciones[voz_elegida_nombre]

            # --- Generar nombre de archivo personalizado ---
            import re
            from datetime import datetime

            # Tomar primeras 5 palabras del texto
            palabras = texto.strip().split()
            if not palabras:
                palabras = ["audio"]
            primeras = palabras[:5]
            # Limpiar cada palabra: solo caracteres alfanuméricos y guiones bajos
            palabras_limpias = [re.sub(r'[^\w]', '', p) for p in primeras]
            # Unir con guiones bajos
            prefijo = "_".join(palabras_limpias) if palabras_limpias else "audio"
            # Limitar longitud para evitar nombres excesivamente largos
            if len(prefijo) > 50:
                prefijo = prefijo[:50]
            # Añadir timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            nombre_archivo = f"{prefijo}_{timestamp}.mp3"
            # ------------------------------------------------

            # Crear archivo temporal
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
                output_file = tmp.name

            # Función asíncrona para generar el audio
            async def generate():
                communicate = edge_tts.Communicate(texto, voice_name)
                await communicate.save(output_file)

            asyncio.run(generate())

            # Leer el archivo generado
            with open(output_file, "rb") as f:
                audio_bytes = f.read()

            # Guardar en session_state
            st.session_state.audio_bytes = audio_bytes
            st.session_state.audio_generado = True
            st.session_state.nombre_archivo = nombre_archivo  # ¡Nuevo nombre!

            # Eliminar archivo temporal
            os.unlink(output_file)

            # Mostrar el reproductor de audio
            st.audio(audio_bytes, format="audio/mp3")
            st.success("¡Audio generado con éxito!")
            

if st.session_state.get("audio_generado"):
    st.download_button(
        label="📥 Descargar audio",
        data=st.session_state.audio_bytes,
        file_name=st.session_state.nombre_archivo,
        mime="audio/mp3"
    )

st.markdown("---")
st.caption("Voces proporcionadas por Microsoft Edge TTS.")
